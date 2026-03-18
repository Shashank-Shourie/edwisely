import requests
from pypdf import PdfReader
from docx import Document
import os
from dotenv import load_dotenv
from huggingface_hub import InferenceClient
from groq import Groq

load_dotenv()

HF_API_KEY = os.getenv("HF_API_KEY")
MODEL = os.getenv("MODEL")


# ------------------------
# 📄 Extract text
# ------------------------

def extract_text_from_pdf(url: str):
    import tempfile
    import requests

    response = requests.get(url)
    with tempfile.NamedTemporaryFile(delete=False) as f:
        f.write(response.content)
        reader = PdfReader(f.name)

        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""

    return text


# def extract_text_from_docx(url: str):
#     import tempfile
#     import requests

#     response = requests.get(url)
#     with tempfile.NamedTemporaryFile(delete=False) as f:
#         f.write(response.content)
#         doc = Document(f.name)

#         return "\n".join([p.text for p in doc.paragraphs])
import tempfile

def extract_text_from_docx(url: str):
    response = requests.get(url)

    if response.status_code != 200:
        raise Exception("Failed to download DOCX")

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as f:
        f.write(response.content)
        f.flush()  # 🔥 IMPORTANT

        doc = Document(f.name)
        return "\n".join([p.text for p in doc.paragraphs])

def extract_text(url: str):
    if url.endswith(".pdf"):
        return extract_text_from_pdf(url)
    else:
        return extract_text_from_docx(url)


# ------------------------
# 🤖 AI Comparison
# ------------------------

# def compare_docs(old_url: str, new_url: str):
#     print(MODEL,"\n")
#     print("Comparing docs\n")

#     try:
#         print("Fetching OLD:", old_url)
#         old_text = extract_text(old_url)[:4000]

#         print("Fetching NEW:", new_url)
#         new_text = extract_text(new_url)[:4000]

#     except Exception as e:
#         print("ERROR in extraction:", str(e))
#         return f"Extraction failed: {str(e)}"

#     prompt = f"""
# Compare the following two syllabus documents and summarize ONLY the differences.

# OLD DOCUMENT:
# {old_text}

# NEW DOCUMENT:
# {new_text}

# Give:
# 1. Added topics
# 2. Removed topics
# 3. Modified topics
# Keep it concise.
# """

#     response = requests.post(
#         f"https://router.huggingface.co/hf-inference/models/{MODEL}",
#         headers={"Authorization": f"Bearer {HF_API_KEY}"},
#         json={
#             "inputs": prompt,
#             "parameters": {
#                 "max_new_tokens": 300,
#                 "temperature": 0.3
#             }
#         }
#     )

#     print("HF STATUS:", response.status_code)
#     print("HF RESPONSE:", response.text)   # 🔥 VERY IMPORTANT

#     try:
#         result = response.json()
#     except:
#         return f"HF Error: {response.text}"

#     # 🔥 Handle all cases
#     if isinstance(result, list):
#         return result[0].get("generated_text", "No output")

#     if "generated_text" in result:
#         return result["generated_text"]

#     if "error" in result:
#         return f"HF Error: {result['error']}"

#     return "AI comparison failed"

# client = InferenceClient(
#     # provider="hf-inference",
#     provider="nebius",
#     api_key=os.getenv("HF_API_KEY"),
# )

client = Groq(api_key=os.getenv("GROQ_API"))

def compare_docs(old_url: str, new_url: str):
    try:
        old_text = extract_text(old_url)[:2000]
        new_text = extract_text(new_url)[:2000]
    except Exception as e:
        return f"Extraction failed: {str(e)}"

    prompt = f"""
Compare the following two syllabus documents and summarize ONLY the differences.

OLD:
{old_text}

NEW:
{new_text}

Return:
- Added topics
- Removed topics
- Modified topics
"""

    try:
        response = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
        )

        return response.choices[0].message.content

    except Exception as e:
        return f"AI Error: {str(e)}"